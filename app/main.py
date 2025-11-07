"""
Elite Pakistani Legal Petition System - FastAPI Backend
Production-grade multi-court petition drafting engine with comprehensive validation.
"""

from __future__ import annotations

import asyncio
import hashlib
import io
import json
import logging
import os
import re
from datetime import datetime
from enum import Enum
from typing import Any, Dict, List, Optional
from uuid import uuid4

from fastapi import BackgroundTasks, FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, Response, StreamingResponse
from pydantic import BaseModel, Field, field_validator, ConfigDict
from dotenv import load_dotenv
from pinecone import Pinecone, ServerlessSpec


logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("elite_petition_ai")


try:
    from docx import Document

    docx_available = True
except ImportError:  # pragma: no cover - optional dependency
    Document = None  # type: ignore
    docx_available = False
    logger.warning(
        "python-docx not installed. DOCX download endpoint will be disabled."
    )

try:
    from bs4 import BeautifulSoup, NavigableString

    bs4_available = True
except ImportError:  # pragma: no cover - optional dependency
    BeautifulSoup = None  # type: ignore
    NavigableString = None  # type: ignore
    bs4_available = False
    logger.warning(
        "beautifulsoup4 not installed. DOCX conversion will fallback to plain text."
    )

try:
    from openai import OpenAI
except ImportError:  # pragma: no cover
    OpenAI = None  # type: ignore


# -----------------------------------------------------------------------------
# Environment Setup
# -----------------------------------------------------------------------------

load_dotenv()


# -----------------------------------------------------------------------------
# External Service Configuration
# -----------------------------------------------------------------------------

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
PINECONE_ENVIRONMENT = os.getenv("PINECONE_ENVIRONMENT", "us-east-1")
PINECONE_INDEX_NAME = os.getenv("PINECONE_INDEX_NAME", "pakistan-legal-corpus")

openai_client: Optional[OpenAI] = None
if not OpenAI:
    logger.warning("openai package not installed; generative features disabled")
elif not OPENAI_API_KEY:
    logger.warning("OPENAI_API_KEY not set; generative features disabled")
else:
    try:
        openai_client = OpenAI(api_key=OPENAI_API_KEY)
        logger.info("OpenAI client initialised")
    except Exception as exc:  # pragma: no cover - defensive
        logger.error("Failed to initialise OpenAI client: %s", exc)
        openai_client = None

pc: Optional[Pinecone] = None
if PINECONE_API_KEY:
    try:
        pc = Pinecone(api_key=PINECONE_API_KEY)
    except Exception as exc:  # pragma: no cover - defensive
        logger.error("Failed to initialise Pinecone client: %s", exc)
        pc = None
else:
    logger.warning("PINECONE_API_KEY not set; vector retrieval disabled")


# -----------------------------------------------------------------------------
# FastAPI Application
# -----------------------------------------------------------------------------

app = FastAPI(
    title="Pakistani Legal Petition AI System - Elite Edition",
    description="Multi-court, citation-verified, comprehensive petition drafting backend",
    version="2.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# -----------------------------------------------------------------------------
# Enumerations & Static Data
# -----------------------------------------------------------------------------


class CourtType(str, Enum):
    SUPREME_COURT = "Supreme Court of Pakistan"
    LAHORE_HIGH_COURT = "Lahore High Court"
    SINDH_HIGH_COURT = "Sindh High Court"
    ISLAMABAD_HIGH_COURT = "Islamabad High Court"
    PESHAWAR_HIGH_COURT = "Peshawar High Court"
    BALOCHISTAN_HIGH_COURT = "Balochistan High Court"
    SESSIONS_COURT = "Sessions Court"
    DISTRICT_COURT = "District Court"


class CaseType(str, Enum):
    CONSTITUTIONAL_PETITION = "constitutional_petition"
    CIVIL_REVISION = "civil_revision"
    CRIMINAL_REVISION = "criminal_revision"
    BAIL_APPLICATION = "bail_application"
    WRIT_PETITION = "writ_petition"
    CRIMINAL_APPEAL = "criminal_appeal"
    CIVIL_APPEAL = "civil_appeal"
    CIVIL_SUIT = "civil_suit"
    CRIMINAL_COMPLAINT = "criminal_complaint"
    ELECTION_PETITION = "election_petition"
    SERVICE_MATTER = "service_matter"
    TAX_MATTER = "tax_matter"


COMPREHENSIVE_TEMPLATES: Dict[str, Dict[str, Any]] = {
    "constitutional_petition": {
        "courts": [
            CourtType.SUPREME_COURT,
            CourtType.LAHORE_HIGH_COURT,
            CourtType.SINDH_HIGH_COURT,
            CourtType.ISLAMABAD_HIGH_COURT,
        ],
        "meta": {
            "case_type": "Constitutional Petition",
            "legal_basis": ["Article 184(3)", "Article 199"],
            "governing_rules": ["Supreme Court Rules 1980", "High Court Rules"],
            "mandatory_requirements": [
                "Vakalatnama",
                "Affidavit",
                "Index",
                "Cause Title",
            ],
        },
        "structure": [
            {
                "id": "cause_title",
                "label": "CAUSE TITLE",
                "required": True,
                "format": "court_specific",
                "content_type": "heading",
            },
            {
                "id": "index",
                "label": "INDEX OF CONTENTS",
                "required": True,
                "format": "table",
                "includes": ["Sections", "Annexures", "Page Numbers"],
            },
            {
                "id": "synopsis",
                "label": "SYNOPSIS",
                "required": True,
                "max_length": 500,
                "format": "numbered_points",
            },
            {
                "id": "list_of_parties",
                "label": "LIST OF PARTIES",
                "required": True,
                "includes": [
                    "Full Names",
                    "CNIC",
                    "Addresses",
                    "Father's Name",
                ],
            },
            {
                "id": "jurisdiction",
                "label": "JURISDICTION",
                "required": True,
                "must_cite": ["Article 184(3)", "Article 199", "Relevant precedents"],
                "min_length": 300,
            },
            {
                "id": "facts",
                "label": "STATEMENT OF FACTS",
                "required": True,
                "format": "chronological_numbered",
                "requirements": [
                    "Dates",
                    "Events",
                    "Documentary references",
                    "Annexure cross-references",
                ],
                "min_length": 1000,
            },
            {
                "id": "grounds",
                "label": "GROUNDS (LEGAL ARGUMENTS)",
                "required": True,
                "format": "numbered_grounds",
                "requirements": [
                    "Constitutional provisions cited",
                    "Case law with exact citations",
                    "Statutory provisions",
                    "Documentary evidence references",
                ],
                "min_length": 1500,
            },
            {
                "id": "interim_relief",
                "label": "INTERIM RELIEF",
                "required": False,
                "requirements": [
                    "Prima facie case",
                    "Irreparable loss",
                    "Balance of convenience",
                ],
            },
            {
                "id": "prayers",
                "label": "PRAYER",
                "required": True,
                "format": "lettered_subparagraphs",
                "requirements": [
                    "Specific relief",
                    "Interim prayer",
                    "Costs",
                    "Any other relief",
                ],
            },
            {
                "id": "annexures",
                "label": "LIST OF ANNEXURES",
                "required": True,
                "format": "table",
                "columns": ["Sr. No.", "Annexure", "Description", "Page No."],
            },
            {
                "id": "verification",
                "label": "VERIFICATION/AFFIDAVIT",
                "required": True,
                "format": "prescribed",
                "must_include": ["Deponent details", "Verification clause", "Date and place"],
            },
            {
                "id": "vakalatnama",
                "label": "VAKALATNAMA",
                "required": True,
                "format": "prescribed_form",
            },
        ],
    },
    "civil_revision": {
        "courts": [
            CourtType.LAHORE_HIGH_COURT,
            CourtType.SINDH_HIGH_COURT,
            CourtType.ISLAMABAD_HIGH_COURT,
        ],
        "meta": {
            "case_type": "Civil Revision",
            "legal_basis": ["Section 115 CPC"],
            "governing_rules": ["CPC 1908", "High Court Rules"],
            "time_limit": "90 days from order",
        },
        "structure": [
            {"id": "cause_title", "label": "CAUSE TITLE", "required": True},
            {"id": "index", "label": "INDEX", "required": True},
            {"id": "synopsis", "label": "SYNOPSIS", "required": True},
            {"id": "parties", "label": "PARTIES", "required": True},
            {
                "id": "jurisdiction",
                "label": "JURISDICTION UNDER SECTION 115 CPC",
                "required": True,
            },
            {
                "id": "facts",
                "label": "FACTS",
                "required": True,
                "min_length": 800,
            },
            {
                "id": "grounds",
                "label": "GROUNDS FOR REVISION",
                "required": True,
                "min_length": 1200,
            },
            {"id": "prayers", "label": "PRAYER", "required": True},
            {"id": "limitation", "label": "LIMITATION", "required": True},
            {"id": "annexures", "label": "ANNEXURES", "required": True},
            {"id": "verification", "label": "VERIFICATION", "required": True},
        ],
    },
    "bail_application": {
        "courts": [
            CourtType.SESSIONS_COURT,
            CourtType.LAHORE_HIGH_COURT,
            CourtType.SINDH_HIGH_COURT,
        ],
        "meta": {
            "case_type": "Bail Application",
            "legal_basis": ["Section 497 CrPC", "Section 498 CrPC", "Article 199"],
            "governing_rules": ["CrPC 1898", "High Court Criminal Rules"],
            "urgency": "High - immediate liberty interest",
        },
        "structure": [
            {"id": "cause_title", "label": "TITLE", "required": True},
            {"id": "case_details", "label": "CASE DETAILS", "required": True},
            {
                "id": "facts",
                "label": "BRIEF FACTS",
                "required": True,
                "min_length": 600,
            },
            {
                "id": "grounds",
                "label": "GROUNDS FOR BAIL",
                "required": True,
                "min_length": 1000,
            },
            {"id": "precedents", "label": "CASE LAW", "required": True},
            {"id": "prayers", "label": "PRAYER", "required": True},
            {"id": "verification", "label": "VERIFICATION", "required": True},
        ],
    },
    "civil_suit": {
        "courts": [CourtType.DISTRICT_COURT, CourtType.SESSIONS_COURT],
        "meta": {
            "case_type": "Civil Suit",
            "legal_basis": ["Order VII CPC"],
            "governing_rules": ["CPC 1908", "Local Court Rules"],
            "court_fee": "As per Court Fees Act",
        },
        "structure": [
            {"id": "heading", "label": "COURT HEADING", "required": True},
            {"id": "parties", "label": "PARTIES", "required": True},
            {"id": "valuation", "label": "VALUATION", "required": True},
            {"id": "jurisdiction", "label": "JURISDICTION", "required": True},
            {
                "id": "facts",
                "label": "FACTS CONSTITUTING CAUSE OF ACTION",
                "required": True,
                "min_length": 1000,
            },
            {"id": "relief", "label": "RELIEF CLAIMED", "required": True},
            {"id": "prayers", "label": "PRAYER", "required": True},
            {"id": "verification", "label": "VERIFICATION", "required": True},
        ],
    },
    "criminal_appeal": {
        "courts": [CourtType.SESSIONS_COURT, CourtType.LAHORE_HIGH_COURT],
        "meta": {
            "case_type": "Criminal Appeal",
            "legal_basis": ["Section 374 CrPC", "Section 417 CrPC"],
            "governing_rules": ["CrPC 1898"],
            "time_limit": "30 days from conviction",
        },
        "structure": [
            {"id": "memo_of_appeal", "label": "MEMORANDUM OF APPEAL", "required": True},
            {"id": "parties", "label": "PARTIES", "required": True},
            {"id": "impugned_order", "label": "IMPUGNED ORDER DETAILS", "required": True},
            {
                "id": "facts",
                "label": "FACTS",
                "required": True,
                "min_length": 1000,
            },
            {
                "id": "grounds",
                "label": "GROUNDS OF APPEAL",
                "required": True,
                "min_length": 1500,
            },
            {"id": "prayers", "label": "PRAYER", "required": True},
            {"id": "limitation", "label": "LIMITATION/CONDONATION", "required": True},
            {"id": "verification", "label": "VERIFICATION", "required": True},
        ],
    },
}


LEGAL_AUTHORITIES: Dict[str, Any] = {
    "constitution": {
        "title": "Constitution of Islamic Republic of Pakistan, 1973",
        "articles": {
            "4": "Right to be dealt with in accordance with law",
            "8": "Laws inconsistent with or in derogation of Fundamental Rights",
            "9": "Security of person",
            "10": "Safeguards as to arrest and detention",
            "10A": "Right to fair trial",
            "14": "Inviolability of dignity of man",
            "15": "Freedom of movement",
            "16": "Freedom of assembly",
            "17": "Freedom of association",
            "18": "Freedom of trade, business or profession",
            "19": "Freedom of speech",
            "19A": "Right to information",
            "23": "Provision as to property",
            "24": "Protection of property rights",
            "25": "Equality of citizens",
            "184": "Original and advisory jurisdiction of Supreme Court",
            "199": "Jurisdiction of High Court",
            "203D": "Powers, jurisdiction and functions of Federal Shariat Court",
        },
    },
    "cpc_1908": {
        "title": "Code of Civil Procedure, 1908",
        "key_provisions": {
            "Order_VII_Rule_1": "Contents of plaint - name of court, parties, facts, cause of action, jurisdiction, relief",
            "Order_VII_Rule_11": "Rejection of plaint - no cause of action, undervalued, barred by law",
            "Section_9": "Courts to try all civil suits unless barred",
            "Section_11": "Res judicata",
            "Section_115": "Revision - exercise of jurisdiction illegally or with material irregularity",
            "Order_VI_Rule_15": "Verification of pleadings",
            "Section_151": "Inherent powers of court",
        },
    },
    "crpc_1898": {
        "title": "Code of Criminal Procedure, 1898",
        "key_provisions": {
            "Section_497": "Bail in non-bailable cases",
            "Section_498": "Bail in High Court or Court of Session",
            "Section_561A": "Inherent powers of High Court",
            "Section_426": "Transfer of cases",
            "Section_439": "Special powers of High Court regarding bail",
        },
    },
    "landmark_cases": {
        "constitutional_law": [
            "District Bar Association, Rawalpindi v. Federation of Pakistan (PLD 2015 SC 401) - Constitutional amendments and basic structure",
            "Asma Jilani v. Government of Punjab (PLD 1972 SC 139) - Fundamental rights and judicial review",
            "State v. Dosso (PLD 1958 SC 533) - Constitutional validity",
            "Mehram Ali v. Federation of Pakistan (PLD 1998 SC 1445) - Article 199 scope",
            "Syed Zafar Ali Shah v. Pervez Musharraf (PLD 2000 SC 869) - Doctrine of necessity",
        ],
        "bail_jurisprudence": [
            "State v. Khizar Hayat (1994 SCMR 1154) - Rule and exception in bail",
            "Ghulam Hussain v. State (PLD 1977 SC 653) - Further inquiry not foreclosed",
            "Muhammad Aslam v. State (2020 SCMR 96) - Bail is rule, jail is exception",
        ],
        "civil_procedure": [
            "Muhammad Sharif v. Waqar Ahmad (2007 SCMR 1450) - Section 115 CPC scope",
            "Messrs Friends Cooperative Housing Society v. DHA (2016 SCMR 1440) - Res judicata",
        ],
    },
}


# -----------------------------------------------------------------------------
# Pydantic Models
# -----------------------------------------------------------------------------


class PetitionRequest(BaseModel):
    case_type: CaseType
    court: CourtType
    parties: Dict[str, Any] = Field(
        ...,
        description="Petitioner and respondent details with CNIC, addresses",
    )
    facts: str = Field(..., min_length=100, description="Detailed chronological facts")
    specific_provisions_challenged: Optional[List[str]] = Field(
        default_factory=list, description="For constitutional petitions"
    )
    relief_sought: str = Field(..., min_length=50)
    urgency_factors: Optional[Dict[str, str]] = Field(default_factory=dict)
    case_details: Optional[Dict[str, str]] = Field(
        default_factory=dict, description="FIR no, case no, court below, etc"
    )
    annexures: List[Dict[str, str]] = Field(
        default_factory=list, description="List of documents to be annexed"
    )

    @field_validator("facts")
    @classmethod
    def validate_facts(cls, value: str) -> str:
        if len(value.split()) < 50:
            raise ValueError("Facts must contain at least 50 words for proper context")
        return value


class CitationSource(BaseModel):
    chunk_id: str
    source_title: str
    source_type: str
    jurisdiction: str
    section_or_article: str
    page_num: Optional[int]
    text_excerpt: str
    similarity_score: float
    url: Optional[str]
    citation_format: str


class ValidationCheck(BaseModel):
    check_name: str
    status: str
    message: str
    severity: str
    fix_suggestion: Optional[str] = None


class PetitionSection(BaseModel):
    id: str
    label: str
    content: str
    word_count: int
    citations: List[str]
    validation_notes: List[str]


class LegacyPartyInfo(BaseModel):
    model_config = ConfigDict(extra="allow")

    petitioner: str
    respondent: str


class LegacyCaseData(BaseModel):
    model_config = ConfigDict(extra="allow")

    case_type: str
    jurisdiction: str
    facts: str
    parties: LegacyPartyInfo
    prayers: Optional[str] = None
    annexures: Optional[List[str]] = Field(default_factory=list)


class LegacyChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None


class LegacyFinalizeRequest(BaseModel):
    approver_name: str
    approver_id: str
    notes: Optional[str] = None


# -----------------------------------------------------------------------------
# Vector Store Manager
# -----------------------------------------------------------------------------


class EnhancedVectorStoreManager:
    def __init__(self) -> None:
        self.index_name = PINECONE_INDEX_NAME
        self.dimension = 1536
        self.index = None
        if pc is None:
            logger.warning("Pinecone unavailable; skipping vector index initialisation")
            return
        self.initialize_index()

    def initialize_index(self) -> None:
        if pc is None:
            return
        try:
            existing_indexes = pc.list_indexes()
            existing_names = [idx.name for idx in existing_indexes]
            if self.index_name not in existing_names:
                pc.create_index(
                    name=self.index_name,
                    dimension=self.dimension,
                    metric="cosine",
                    spec=ServerlessSpec(cloud="aws", region=PINECONE_ENVIRONMENT),
                )
            self.index = pc.Index(self.index_name)
            logger.info("Connected to Pinecone index: %s", self.index_name)
        except Exception as exc:
            logger.error("Error initialising Pinecone: %s", exc)
            self.index = None

    async def upsert_texts(
        self,
        items: List[Dict[str, Any]],
        namespace: Optional[str] = None,
        batch_size: int = 32,
    ) -> None:
        """Embed and upsert textual chunks into Pinecone."""

        if not items:
            logger.info("No items provided for upsert; skipping")
            return

        if self.index is None:
            raise RuntimeError("Pinecone index is not initialised")

        if not openai_client:
            raise RuntimeError("OpenAI client unavailable for embedding generation")

        ns = namespace or "knowledge-base"

        for start in range(0, len(items), batch_size):
            batch = items[start : start + batch_size]
            texts = [item.get("text", "") for item in batch]

            embeddings = await asyncio.gather(
                *[self.generate_embedding(text) for text in texts]
            )

            vectors = []
            for item, embedding in zip(batch, embeddings):
                metadata = item.get("metadata", {}) or {}
                metadata.setdefault("source_title", item.get("source_title"))
                metadata.setdefault("doc_type", item.get("doc_type", "knowledge_base"))
                metadata.setdefault("jurisdiction", item.get("jurisdiction", "Pakistan"))
                vectors.append(
                    {
                        "id": item.get("id") or str(uuid4()),
                        "values": embedding,
                        "metadata": metadata,
                    }
                )

            try:
                self.index.upsert(vectors=vectors, namespace=ns)
                logger.info(
                    "Upserted %s vectors to namespace '%s'", len(vectors), ns
                )
            except Exception as exc:
                logger.error("Failed to upsert batch to Pinecone: %s", exc)
                raise

    async def search_with_filters(
        self, query: str, case_type: str, court: str, top_k: int = 15
    ) -> List[Dict[str, Any]]:
        if self.index is None or not openai_client:
            return []

        try:
            query_embedding = await self.generate_embedding(query)
        except Exception:
            return []

        filter_dict = {
            "$or": [
                {"jurisdiction": court},
                {"jurisdiction": "Pakistan"},
                {"jurisdiction": "All Courts"},
            ],
            "relevance_to_case_type": {"$in": [case_type, "general"]},
        }

        try:
            results = self.index.query(
                vector=query_embedding,
                top_k=top_k,
                include_metadata=True,
                filter=filter_dict,
            )
        except Exception as exc:
            logger.error("Error querying Pinecone: %s", exc)
            return []

        formatted_results: List[Dict[str, Any]] = []
        for match in getattr(results, "matches", []) or []:
            metadata = match.metadata or {}
            formatted_results.append(
                {
                    "id": match.id,
                    "score": match.score,
                    "metadata": metadata,
                    "text": metadata.get("text", ""),
                }
            )
        return formatted_results

    async def generate_embedding(self, text: str) -> List[float]:
        if not openai_client:
            raise RuntimeError("OpenAI client unavailable")

        response = await asyncio.to_thread(
            openai_client.embeddings.create,
            model="text-embedding-3-small",
            input=text,
        )
        return response.data[0].embedding


vector_store: Optional[EnhancedVectorStoreManager] = None
try:
    vector_store = EnhancedVectorStoreManager()
except Exception as exc:  # pragma: no cover
    logger.error("Vector store initialisation failed: %s", exc)
    vector_store = None


# -----------------------------------------------------------------------------
# Petition Generator
# -----------------------------------------------------------------------------


class ElitePetitionGenerator:
    @staticmethod
    async def generate_complete_petition(request: PetitionRequest) -> Dict[str, Any]:
        template = COMPREHENSIVE_TEMPLATES.get(request.case_type.value)
        if not template:
            raise ValueError(f"Template not found for case type: {request.case_type.value}")

        if request.court not in template["courts"]:
            raise ValueError(f"{request.case_type.value} not available in {request.court.value}")

        sources = await ElitePetitionGenerator.retrieve_legal_sources(request=request, top_k=20)

        sections: List[PetitionSection] = []
        for section_spec in template["structure"]:
            section = await ElitePetitionGenerator.generate_section(
                section_spec=section_spec,
                request=request,
                sources=sources,
                template=template,
            )
            sections.append(section)

        validation = await ElitePetitionGenerator.validate_petition(
            sections=sections, template=template, request=request
        )

        provenance = ElitePetitionGenerator.build_comprehensive_provenance(sources)

        draft_id = hashlib.sha256(
            f"{request.facts}{request.court.value}{datetime.now().isoformat()}".encode()
        ).hexdigest()[:16]

        return {
            "draft_id": draft_id,
            "case_type": request.case_type.value,
            "court": request.court.value,
            "template_version": template["meta"].get("case_type", "v1"),
            "meta": template["meta"],
            "sections": [section.dict() for section in sections],
            "annexures": request.annexures,
            "provenance": [source.dict() for source in provenance],
            "validation": validation,
            "created_at": datetime.now().isoformat(),
            "total_word_count": sum(section.word_count for section in sections),
            "citation_count": len(provenance),
        }

    @staticmethod
    async def retrieve_legal_sources(request: PetitionRequest, top_k: int) -> List[Dict[str, Any]]:
        query_components = [
            request.facts,
            request.relief_sought,
            request.case_type.value,
            request.court.value,
        ]
        query_components.extend(request.specific_provisions_challenged or [])
        combined_query = " ".join(query_components)

        if vector_store:
            results = await vector_store.search_with_filters(
                query=combined_query,
                case_type=request.case_type.value,
                court=request.court.value,
                top_k=top_k,
            )
            if results:
                return results

        return ElitePetitionGenerator.generate_fallback_sources(request, top_k=top_k)

    @staticmethod
    def generate_fallback_sources(
        request: PetitionRequest, top_k: int = 20
    ) -> List[Dict[str, Any]]:
        fallback_sources: List[Dict[str, Any]] = []

        constitution = LEGAL_AUTHORITIES.get("constitution", {})
        for article, summary in constitution.get("articles", {}).items():
            fallback_sources.append(
                {
                    "id": f"constitution-{article}",
                    "score": 0.75,
                    "metadata": {
                        "source_title": constitution.get("title", "Constitution"),
                        "doc_type": "constitution",
                        "jurisdiction": "Pakistan",
                        "section": f"Article {article}",
                        "citation": f"Article {article} Constitution of Pakistan",
                    },
                    "text": summary,
                }
            )

        case_bucket = LEGAL_AUTHORITIES.get("landmark_cases", {})
        for topic, cases in case_bucket.items():
            for idx, case in enumerate(cases):
                fallback_sources.append(
                    {
                        "id": f"case-{topic}-{idx}",
                        "score": 0.7,
                        "metadata": {
                            "source_title": topic.replace("_", " ").title(),
                            "doc_type": "case_law",
                            "jurisdiction": "Pakistan",
                            "citation": case.split(" - ")[0],
                        },
                        "text": case,
                    }
                )
        limit = top_k if top_k and top_k > 0 else 20
        return fallback_sources[:limit]

    @staticmethod
    async def generate_section(
        section_spec: Dict[str, Any],
        request: PetitionRequest,
        sources: List[Dict[str, Any]],
        template: Dict[str, Any],
    ) -> PetitionSection:
        section_id = section_spec["id"]

        if section_id == "cause_title":
            content = ElitePetitionGenerator.generate_cause_title(request, template)
        elif section_id == "index":
            content = ElitePetitionGenerator.generate_index(template)
        elif section_id in {"list_of_parties", "parties"}:
            content = ElitePetitionGenerator.generate_parties_table(request, section_spec)
        elif section_id == "jurisdiction":
            content = await ElitePetitionGenerator.generate_jurisdiction_section(request, sources)
        elif section_id == "facts":
            content = await ElitePetitionGenerator.generate_facts_section(request, sources)
        elif section_id == "grounds":
            content = await ElitePetitionGenerator.generate_grounds_section(request, sources)
        elif section_id == "prayers":
            content = await ElitePetitionGenerator.generate_prayers_section(request, sources)
        elif section_id == "verification":
            content = ElitePetitionGenerator.generate_verification(request)
        elif section_id == "annexures":
            content = ElitePetitionGenerator.generate_annexures_list(request)
        elif section_id == "vakalatnama":
            content = ElitePetitionGenerator.generate_vakalatnama_note(request)
        else:
            content = await ElitePetitionGenerator.generate_generic_section(
                section_id=section_id,
                section_spec=section_spec,
                request=request,
                sources=sources,
            )

        word_count = len(re.findall(r"\w+", content))
        citations = re.findall(r"\[(.*?)\]", content)

        return PetitionSection(
            id=section_id,
            label=section_spec["label"],
            content=content,
            word_count=word_count,
            citations=citations,
            validation_notes=[],
        )

    @staticmethod
    def generate_cause_title(request: PetitionRequest, template: Dict[str, Any]) -> str:
        petitioner = request.parties.get("petitioner", {}).get("name", "[PETITIONER NAME]")
        respondent = request.parties.get("respondent", {}).get("name", "[RESPONDENT NAME]")

        if request.court == CourtType.SUPREME_COURT:
            return f"""
<div class="text-center font-bold mb-6">
    <h1 class="text-xl">IN THE SUPREME COURT OF PAKISTAN</h1>
    <h2 class="text-lg mt-2">(ORIGINAL/APPELLATE JURISDICTION)</h2>
    <br/>
    <h3>{template['meta']['case_type'].upper()} NO. _______ OF {datetime.now().year}</h3>
    <br/>
    <div class="mt-4">
        <p>{petitioner.upper()}</p>
        <p class="text-sm">(Petitioner)</p>
        <br/>
        <p class="font-bold">VERSUS</p>
        <br/>
        <p>{respondent.upper()}</p>
        <p class="text-sm">(Respondent)</p>
    </div>
</div>
"""

        if "High Court" in request.court.value:
            return f"""
<div class="text-center font-bold mb-6">
    <h1 class="text-xl">IN THE HONOURABLE {request.court.value.upper()}</h1>
    <h2 class="text-lg mt-2">({template['meta']['case_type'].upper()})</h2>
    <br/>
    <h3>{template['meta']['case_type'].upper()} NO. _______ OF {datetime.now().year}</h3>
    <br/>
    <table class="mx-auto mt-4" style="border-collapse: collapse;">
        <tr>
            <td class="text-left pr-8">{petitioner.upper()}</td>
            <td class="text-right">PETITIONER</td>
        </tr>
        <tr>
            <td colspan="2" class="text-center py-2 font-bold">VERSUS</td>
        </tr>
        <tr>
            <td class="text-left pr-8">{respondent.upper()}</td>
            <td class="text-right">RESPONDENT</td>
        </tr>
    </table>
</div>
"""

        return f"""
<div class="text-center font-bold mb-6">
    <h1>IN THE COURT OF {request.court.value.upper()}</h1>
    <h2 class="mt-2">{template['meta']['case_type'].upper()}</h2>
    <br/>
    <p class="mt-4">{petitioner} vs {respondent}</p>
</div>
"""

    @staticmethod
    def generate_index(template: Dict[str, Any]) -> str:
        rows = "\n".join(
            [
                f"<tr><td>{idx + 1}</td><td>{section['label']}</td><td>Page __</td></tr>"
                for idx, section in enumerate(template["structure"])
            ]
        )
        return f"""
<div class="mb-8">
    <h2 class="text-center font-bold text-lg mb-4">INDEX OF CONTENTS</h2>
    <table class="w-full border-collapse border-2 border-gray-800">
        <thead>
            <tr class="bg-gray-200">
                <th class="border border-gray-800 p-2 w-16">Sr.No.</th>
                <th class="border border-gray-800 p-2">Particulars</th>
                <th class="border border-gray-800 p-2 w-24">Page No.</th>
            </tr>
        </thead>
        <tbody>
            {rows}
        </tbody>
    </table>
</div>
"""

    @staticmethod
    def generate_parties_table(request: PetitionRequest, section_spec: Dict[str, Any]) -> str:
        petitioner = request.parties.get("petitioner", {})
        respondent = request.parties.get("respondent", {})

        def render_party(label: str, party: Dict[str, Any]) -> str:
            return """
            <tr>
                <td class="border border-gray-800 p-2 font-bold">{label}</td>
                <td class="border border-gray-800 p-2">{name}</td>
                <td class="border border-gray-800 p-2">{cnic}</td>
                <td class="border border-gray-800 p-2">{address}</td>
            </tr>
            """.format(
                label=label,
                name=party.get("name", "[Full Name]"),
                cnic=party.get("cnic", "[CNIC]") or "[CNIC]",
                address=party.get("address", "[Address]") or "[Address]",
            )

        return f"""
<div class="mb-8">
    <h2 class="font-bold text-lg mb-4 underline">{section_spec['label']}</h2>
    <table class="w-full border-collapse border-2 border-gray-800">
        <thead>
            <tr class="bg-gray-200">
                <th class="border border-gray-800 p-2 w-32">Role</th>
                <th class="border border-gray-800 p-2">Name</th>
                <th class="border border-gray-800 p-2 w-40">CNIC / ID</th>
                <th class="border border-gray-800 p-2">Address</th>
            </tr>
        </thead>
        <tbody>
            {render_party('Petitioner', petitioner)}
            {render_party('Respondent', respondent)}
        </tbody>
    </table>
</div>
"""

    @staticmethod
    async def generate_jurisdiction_section(
        request: PetitionRequest, sources: List[Dict[str, Any]]
    ) -> str:
        jurisdictional_sources = [
            src
            for src in sources
            if any(
                key in (src.get("metadata", {}).get("section", "").lower() or "")
                for key in ["article", "jurisdiction", "section 115", "section 497"]
            )
        ][:5]

        sources_context = "\n\n".join(
            [
                f"[SOURCE {idx + 1}] {src['metadata'].get('source_title', 'Unknown')}\n"
                f"{src['metadata'].get('section', 'N/A')}: {src.get('text', '')[:400]}..."
                for idx, src in enumerate(jurisdictional_sources)
            ]
        )

        prompt = f"""Generate a comprehensive JURISDICTION section for a {request.case_type.value} in {request.court.value}.

Instructions:
1. Cite the exact constitutional/statutory provision conferring jurisdiction.
2. Reference mandatory precedents from the provided sources.
3. Address any jurisdictional bars or alternate remedies.
4. Minimum 300 words, formal legal language.

LEGAL SOURCES:
{sources_context or 'No retrieved sources available. Rely on constitutional provisions.'}

CASE DETAILS:
Court: {request.court.value}
Relief: {request.relief_sought}

Use inline citations in the form [Source Title: Section/Citation]."""

        if not openai_client:
            logger.warning("OpenAI unavailable; returning jurisdiction placeholder")
            return "<div><h2>JURISDICTION</h2><p>[Manual drafting required - OpenAI unavailable]</p></div>"

        try:
            response = await asyncio.to_thread(
                openai_client.chat.completions.create,
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are an expert Pakistani legal drafter. Cite statutory"\
                            " text verbatim, use precise legal terminology, and align"\
                            " with superior court jurisprudence."
                        ),
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.2,
                max_tokens=1500,
            )
            content = response.choices[0].message.content
            return f"""
<div class="mb-8">
    <h2 class="font-bold text-lg mb-4 underline">JURISDICTION</h2>
    <div class="text-justify leading-relaxed">
        {content}
    </div>
</div>
"""
        except Exception as exc:
            logger.error("Error generating jurisdiction section: %s", exc)
            return "<div><h2>JURISDICTION</h2><p>[Error - requires manual drafting]</p></div>"

    @staticmethod
    async def generate_facts_section(
        request: PetitionRequest, sources: List[Dict[str, Any]]
    ) -> str:
        facts_paragraphs = [
            paragraph.strip()
            for paragraph in request.facts.split("\n")
            if paragraph.strip()
        ]

        prompt = f"""Transform the provided factual narrative into a chronological STATEMENT OF FACTS.

Requirements:
1. Number each paragraph (1., 2., 3., ...).
2. Ensure chronological order with explicit dates.
3. Reference annexures where relevant in the format [Annexure-A].
4. Link facts to documents where possible.
5. Minimum 800 words using formal legal tone.

Facts Provided:
{chr(10).join(facts_paragraphs)}

Context:
Court: {request.court.value}
Case Type: {request.case_type.value}
Relief: {request.relief_sought}
Petitioner: {request.parties.get('petitioner', {}).get('name', 'Petitioner')}
Respondent: {request.parties.get('respondent', {}).get('name', 'Respondent')}
"""

        if not openai_client:
            logger.warning("OpenAI unavailable; returning facts placeholder")
            return "<div><h2>STATEMENT OF FACTS</h2><p>[Manual drafting required - OpenAI unavailable]</p></div>"

        try:
            response = await asyncio.to_thread(
                openai_client.chat.completions.create,
                model="gpt-4o",
            messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are a senior Pakistani advocate drafting factual"\
                            " pleadings. Maintain strict chronology, cite dates, and"\
                            " reference annexures and documents."
                        ),
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.25,
                max_tokens=2600,
            )
            content = response.choices[0].message.content
            return f"""
<div class="mb-8">
    <h2 class="font-bold text-lg mb-4 underline">STATEMENT OF FACTS</h2>
    <div class="text-justify leading-relaxed space-y-3">
        {content}
    </div>
</div>
"""
        except Exception as exc:
            logger.error("Error generating facts section: %s", exc)
            return "<div><h2>STATEMENT OF FACTS</h2><p>[Error - requires manual drafting]</p></div>"

    @staticmethod
    async def generate_grounds_section(
        request: PetitionRequest, sources: List[Dict[str, Any]]
    ) -> str:
        legal_sources = [src for src in sources if float(src.get("score", 0)) >= 0.7][:10]
        sources_context = "\n\n".join(
            [
                f"[SOURCE {idx + 1}]\nTitle: {src['metadata'].get('source_title', 'Unknown')}\n"
                f"Section: {src['metadata'].get('section', 'N/A')}\n"
                f"Type: {src['metadata'].get('doc_type', 'N/A')}\n"
                f"Text: {src.get('text', '')[:500]}...\n"
                f"Citation: {src['metadata'].get('citation', 'N/A')}"
                for idx, src in enumerate(legal_sources)
            ]
        )

        prompt = f"""Produce detailed numbered LEGAL GROUNDS (Ground I, Ground II, etc.).

Each ground must:
• State the legal principle.
• Quote the precise constitutional/statutory provision.
• Cite case law with accurate citations (PLD, SCMR, etc.).
• Apply the principle to the provided facts.
• Minimum total length: 1500 words.

LEGAL SOURCES:
{sources_context or 'Rely on core Pakistani legal principles and landmark jurisprudence.'}

FACT SUMMARY:
{request.facts[:600]}

Relief: {request.relief_sought}
Challenged Provisions: {', '.join(request.specific_provisions_challenged) if request.specific_provisions_challenged else 'N/A'}
"""

        if not openai_client:
            logger.warning("OpenAI unavailable; returning grounds placeholder")
            return "<div><h2>GROUNDS</h2><p>[Manual drafting required - OpenAI unavailable]</p></div>"

        try:
            response = await asyncio.to_thread(
                openai_client.chat.completions.create,
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
            "content": (
                            "You are a senior advocate drafting Pakistani petitions."\
                            " Provide rigorous legal grounds with precise citations"\
                            " and detailed application to facts."
                        ),
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.2,
                max_tokens=3000,
            )
            content = response.choices[0].message.content
            return f"""
<div class="mb-8">
    <h2 class="font-bold text-lg mb-4 underline">GROUNDS</h2>
    <div class="text-justify leading-relaxed space-y-4">
        {content}
    </div>
</div>
"""
        except Exception as exc:
            logger.error("Error generating grounds section: %s", exc)
            return "<div><h2>GROUNDS</h2><p>[Error - requires manual drafting]</p></div>"

    @staticmethod
    async def generate_prayers_section(
        request: PetitionRequest, sources: List[Dict[str, Any]]
    ) -> str:
        prompt = f"""Draft the PRAYER section with lettered sub-paragraphs (a), (b), (c)...

Include:
• Interim relief (if appropriate).
• Final substantive relief specific to the case.
• Costs.
• Any other relief deemed just.

Use formal Pakistani legal terminology.

Relief sought summary:
{request.relief_sought}

Case type: {request.case_type.value}
Court: {request.court.value}
"""

        if not openai_client:
            logger.warning("OpenAI unavailable; returning prayer placeholder")
            return "<div><h2>PRAYER</h2><p>[Manual drafting required - OpenAI unavailable]</p></div>"

        try:
            response = await asyncio.to_thread(
                openai_client.chat.completions.create,
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
            "content": (
                            "You are drafting the prayer in a Pakistani petition."\
                            " Be concise, precise, and aligned with local practice."
                        ),
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.15,
                max_tokens=900,
            )
            content = response.choices[0].message.content
            return f"""
<div class="mb-8">
    <h2 class="font-bold text-lg mb-4 underline">PRAYER</h2>
    <div class="leading-relaxed">
        <p class="mb-4">In the circumstances stated, it is respectfully prayed that this Honourable Court may graciously be pleased to:</p>
        {content}
    </div>
</div>
"""
        except Exception as exc:
            logger.error("Error generating prayer section: %s", exc)
            return "<div><h2>PRAYER</h2><p>[Error - requires manual drafting]</p></div>"

    @staticmethod
    def generate_verification(request: PetitionRequest) -> str:
        petitioner = request.parties.get("petitioner", {})
        name = petitioner.get("name", "[NAME]")
        father_name = petitioner.get("father_name", "[FATHER'S NAME]")
        address = petitioner.get("address", "[ADDRESS]")
        return f"""
<div class="mb-8 border-t-2 pt-6 mt-8">
    <h2 class="font-bold text-lg mb-4 underline">VERIFICATION</h2>
    <div class="leading-relaxed text-justify">
        <p class="mb-4">
            I, <strong>{name}</strong>, son/daughter/wife of <strong>{father_name}</strong>,
            resident of <strong>{address}</strong>, the petitioner above-named, do hereby solemnly
            affirm and verify that the contents of the above petition are true and correct to the
            best of my knowledge and belief and nothing material has been concealed therefrom.
        </p>
        <p class="mb-4">
            Verified that the statements contained in paragraphs 1 to ___ are true to my personal
            knowledge and those in paragraphs ___ to ___ are believed to be true on the basis of
            information and documents available with me.
        </p>
        <div class="mt-8 text-right">
            <p class="mb-12">____________________</p>
            <p class="font-bold">{name.upper()}</p>
            <p>Petitioner</p>
        </div>
        <p class="mt-6">
            Verified at <strong>___________</strong> on this <strong>_____</strong> day of
            <strong>___________</strong>, {datetime.now().year}.
        </p>
    </div>
</div>
"""

    @staticmethod
    def generate_annexures_list(request: PetitionRequest) -> str:
        if request.annexures:
            rows = "".join(
                [
                    """
                    <tr>
                        <td class="border border-gray-800 p-2 text-center">{idx}</td>
                        <td class="border border-gray-800 p-2 text-center">Annexure-{letter}</td>
                        <td class="border border-gray-800 p-2">{description}</td>
                        <td class="border border-gray-800 p-2 text-center">Page __</td>
                    </tr>
                    """.format(
                        idx=index + 1,
                        letter=chr(65 + index),
                        description=annex.get("description", "Document"),
                    )
                    for index, annex in enumerate(request.annexures)
                ]
            )
        else:
            rows = """
            <tr>
                <td class="border border-gray-800 p-2 text-center">1</td>
                <td class="border border-gray-800 p-2 text-center">Annexure-A</td>
                <td class="border border-gray-800 p-2">Copy of the impugned order/document</td>
                <td class="border border-gray-800 p-2 text-center">Page __</td>
            </tr>
            """

        return f"""
<div class="mb-8">
    <h2 class="font-bold text-lg mb-4 underline">LIST OF ANNEXURES</h2>
    <table class="w-full border-collapse border-2 border-gray-800">
        <thead>
            <tr class="bg-gray-200">
                <th class="border border-gray-800 p-2 w-20">Sr. No.</th>
                <th class="border border-gray-800 p-2 w-32">Annexure</th>
                <th class="border border-gray-800 p-2">Description</th>
                <th class="border border-gray-800 p-2 w-24">Page No.</th>
            </tr>
        </thead>
        <tbody>
            {rows}
        </tbody>
    </table>
</div>
"""

    @staticmethod
    def generate_vakalatnama_note(request: PetitionRequest) -> str:
        counsel = request.parties.get("counsel", {})
        counsel_name = counsel.get("name", "[COUNSEL NAME]")
        return f"""
<div class="mb-8">
    <h2 class="font-bold text-lg mb-4 underline">VAKALATNAMA</h2>
    <p class="leading-relaxed text-justify">
        A duly executed Vakalatnama in favour of <strong>{counsel_name}</strong>, Advocate High Court,
        is annexed hereto in the prescribed format, authorising representation before this Honourable Court.
    </p>
    <p class="mt-4">
        Counsel Contact: {counsel.get('contact', '[CONTACT NUMBER]')}<br/>
        Enrollment No.: {counsel.get('enrollment_no', '[BAR ENROLLMENT]')}
    </p>
</div>
"""

    @staticmethod
    async def generate_generic_section(
        section_id: str,
        section_spec: Dict[str, Any],
        request: PetitionRequest,
        sources: List[Dict[str, Any]],
    ) -> str:
        instructions: List[str] = []
        for key in ["format", "requirements", "includes", "columns"]:
            value = section_spec.get(key)
            if value:
                instructions.append(f"{key}: {value}")

        source_snippets = "\n\n".join(
            [
                f"{src['metadata'].get('source_title', 'Unknown')} ({src['metadata'].get('citation', 'N/A')}): {src.get('text', '')[:300]}"
                for src in sources[:5]
            ]
        )

        prompt = f"""Draft the section titled "{section_spec['label']}" for a {request.case_type.value} pending before {request.court.value}.

Section requirements:
{chr(10).join(instructions) if instructions else 'Follow Pakistani legal drafting standards for this section.'}

Facts summary:
{request.facts[:500]}

Relief sought:
{request.relief_sought}

Available sources:
{source_snippets or 'Use standard Pakistani legal practice guidance.'}

Return HTML-ready content (no <html> or <body> tags)."""

        if not openai_client:
            logger.warning("OpenAI unavailable; returning placeholder for %s", section_id)
            return f"<div><h2>{section_spec['label']}</h2><p>[Manual drafting required for {section_spec['label']}]</p></div>"

        try:
            response = await asyncio.to_thread(
                openai_client.chat.completions.create,
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
            "content": (
                            "You are producing a section of a Pakistani legal petition."\
                            " Ensure compliance with procedural requirements and"\
                            " integrate citations where appropriate."
                        ),
                    },
                    {"role": "user", "content": prompt},
                ],
                temperature=0.3,
                max_tokens=1500,
            )
            content = response.choices[0].message.content
            return f"""
<div class="mb-8">
    <h2 class="font-bold text-lg mb-4 underline">{section_spec['label']}</h2>
    <div class="leading-relaxed text-justify">
        {content}
    </div>
</div>
"""
        except Exception as exc:
            logger.error("Error generating %s section: %s", section_id, exc)
            return f"<div><h2>{section_spec['label']}</h2><p>[Error - requires manual drafting]</p></div>"

    @staticmethod
    async def validate_petition(
        sections: List[PetitionSection],
        template: Dict[str, Any],
        request: PetitionRequest,
    ) -> Dict[str, Any]:
        section_map = {section.id: section for section in sections}
        checks: List[ValidationCheck] = []
        required_specs = [spec for spec in template["structure"] if spec.get("required")]
        missing_required = 0

        for spec in template["structure"]:
            section = section_map.get(spec["id"])
            if not section or not section.content.strip():
                if spec.get("required"):
                    missing_required += 1
                    checks.append(
            ValidationCheck(
                            check_name=f"section_{spec['id']}",
                            status="missing",
                            severity="critical",
                            message=f"{spec['label']} is required but missing",
                            fix_suggestion="Regenerate section or provide manual draft",
                        )
                    )
                continue

            min_length = spec.get("min_length")
            if min_length and section.word_count < min_length:
                checks.append(
            ValidationCheck(
                        check_name=f"length_{spec['id']}",
                        status="too_short",
                        severity="warning",
                        message=f"{spec['label']} should be at least {min_length} words",
                        fix_suggestion="Expand section with additional detail and authorities",
                    )
                )

            if spec.get("must_cite") and not section.citations:
                checks.append(
            ValidationCheck(
                        check_name=f"citations_{spec['id']}",
                        status="missing_citations",
                        severity="warning",
                        message=f"{spec['label']} requires explicit citations",
                        fix_suggestion="Insert pinpoint references to Articles/Sections or case law",
                    )
                )

        citations_total = sum(len(section.citations) for section in sections)
        if citations_total < 10:
            checks.append(
            ValidationCheck(
                    check_name="citation_density",
                    status="insufficient",
                    severity="warning",
                    message="Fewer than 10 citations detected; consider adding more authorities",
                    fix_suggestion="Reference additional constitutional provisions or precedents",
                )
            )

        coverage_ratio = (
            (len(required_specs) - missing_required) / len(required_specs)
            if required_specs
            else 1.0
        )
        has_critical = any(check.severity == "critical" for check in checks)
        overall_score = max(0.0, min(1.0, 0.6 + 0.35 * coverage_ratio - 0.15 * has_critical))

        status = "pass"
        if has_critical:
            status = "fail"
        elif overall_score < 0.8:
            status = "warn"

        checks.append(
            ValidationCheck(
                check_name="overall_compliance",
                status=status,
                severity="info" if status == "pass" else ("warning" if status == "warn" else "critical"),
                message=f"Overall compliance score: {overall_score:.2f}",
            )
        )

        return {
            "status": status,
            "score": round(overall_score, 2),
            "checks": [check.dict() for check in checks],
        }

    @staticmethod
    def build_comprehensive_provenance(
        sources: List[Dict[str, Any]]
    ) -> List[CitationSource]:
        provenance: List[CitationSource] = []
        for src in sources:
            metadata = src.get("metadata", {})
            provenance.append(
                CitationSource(
                    chunk_id=src.get("id", ""),
                    source_title=metadata.get("source_title")
                    or metadata.get("title")
                    or "Unknown Source",
                    source_type=metadata.get("doc_type", "unknown"),
                    jurisdiction=metadata.get("jurisdiction", "Pakistan"),
                    section_or_article=metadata.get("section")
                    or metadata.get("article")
                    or "N/A",
                    page_num=metadata.get("page") or metadata.get("page_number"),
                    text_excerpt=src.get("text", "")[:400],
                    similarity_score=float(src.get("score", 0.0)),
                    url=metadata.get("url"),
                    citation_format=metadata.get("citation", metadata.get("reference", "")),
                )
            )
        return provenance


# -----------------------------------------------------------------------------
# In-memory Draft Store
# -----------------------------------------------------------------------------


draft_registry: Dict[str, Dict[str, Any]] = {}


# -----------------------------------------------------------------------------
# Helper Functions
# -----------------------------------------------------------------------------


async def ensure_services_ready() -> None:
    if not openai_client:
        raise HTTPException(status_code=503, detail="OpenAI client not configured")


def render_petition_html(draft: Dict[str, Any]) -> str:
    body_sections = "".join(
        [
            f"<div class='section'><h2>{section['label']}</h2>{section['content']}</div>"
            for section in draft.get("sections", [])
        ]
    )
    return f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Legal Petition - {draft.get('draft_id')}</title>
    <style>
        body {{ font-family: 'Times New Roman', serif; margin: 40px; line-height: 1.6; }}
        h1 {{ text-align: center; margin-bottom: 30px; }}
        .section {{ margin-bottom: 28px; }}
        .metadata {{ margin-bottom: 24px; }}
    </style>
</head>
<body>
    <h1>Legal Petition Draft</h1>
    <div class="metadata">
        <p><strong>Draft ID:</strong> {draft.get('draft_id')}</p>
        <p><strong>Court:</strong> {draft.get('court')}</p>
        <p><strong>Case Type:</strong> {draft.get('case_type')}</p>
        <p><strong>Created:</strong> {draft.get('created_at')}</p>
    </div>
    {body_sections}
</body>
</html>
"""


def process_html_element(doc: Document, element, paragraph=None, formatting=None):
    if formatting is None:
        formatting = {}

    if bs4_available and NavigableString and isinstance(element, NavigableString):
        text = str(element)
        if not text:
            return paragraph
        if not text.strip():
            if "\n" in text:
                paragraph = paragraph or doc.add_paragraph()
                paragraph.add_run("\n")
            return paragraph
        paragraph = paragraph or doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = formatting.get("bold", False)
        run.italic = formatting.get("italic", False)
        return paragraph

    if not bs4_available or element is None:
        return paragraph

    name = element.name

    if name == "br":
        paragraph = paragraph or doc.add_paragraph()
        paragraph.add_run("\n")
        return paragraph

    if name in {"strong", "b"}:
        new_format = {**formatting, "bold": True}
        for child in element.children:
            paragraph = process_html_element(doc, child, paragraph, new_format)
        return paragraph

    if name in {"em", "i"}:
        new_format = {**formatting, "italic": True}
        for child in element.children:
            paragraph = process_html_element(doc, child, paragraph, new_format)
        return paragraph

    if name in {"h1", "h2", "h3", "h4"}:
        level_map = {"h1": 0, "h2": 1, "h3": 2, "h4": 3}
        text = element.get_text(strip=True)
        if text:
            doc.add_heading(text, level=level_map.get(name, 1))
        return None

    if name == "p":
        paragraph = doc.add_paragraph()
        for child in element.children:
            paragraph = process_html_element(doc, child, paragraph, formatting)
        return None

    if name == "div" or name == "span":
        for child in element.children:
            paragraph = process_html_element(doc, child, paragraph, formatting)
        return paragraph

    if name in {"ul", "ol"}:
        list_style = "List Bullet" if name == "ul" else "List Number"
        for li in element.find_all("li", recursive=False):
            paragraph = doc.add_paragraph(style=list_style)
            for child in li.children:
                paragraph = process_html_element(doc, child, paragraph, formatting)
        return None

    if name == "li":
        paragraph = doc.add_paragraph(style="List Bullet")
        for child in element.children:
            paragraph = process_html_element(doc, child, paragraph, formatting)
        return None

    if name == "table":
        rows = element.find_all("tr")
        if rows:
            cols = max(len(row.find_all(["th", "td"])) for row in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(rows):
                cells = row.find_all(["th", "td"])
                for c_idx in range(cols):
                    cell = table.cell(r_idx, c_idx)
                    if c_idx < len(cells):
                        cell_text = cells[c_idx].get_text(separator=" ", strip=True)
                    else:
                        cell_text = ""
                    cell.text = cell_text
        return None

    for child in getattr(element, "children", []):
        paragraph = process_html_element(doc, child, paragraph, formatting)
    return paragraph


def render_html_to_docx(doc: Document, html: str) -> None:
    if not html:
        return

    if bs4_available and BeautifulSoup:
        soup = BeautifulSoup(html, "html.parser")
        for child in soup.body.children if soup.body else soup.children:
            process_html_element(doc, child)
    else:
        text = re.sub(r"<br\s*/?>", "\n", html)
        text = re.sub(r"<[^>]+>", "", text)
        for line in text.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line)


def build_petition_docx(draft: Dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading("Pakistani Legal Petition", level=0)

    doc.add_paragraph().add_run("Draft ID: ").bold = True
    doc.paragraphs[-1].add_run(str(draft.get("draft_id") or "N/A"))

    metadata_pairs = [
        ("Court", draft.get("court")),
        ("Case Type", draft.get("case_type")),
        ("Created", draft.get("created_at")),
        ("Template", draft.get("template_version")),
        ("Total Word Count", str(draft.get("total_word_count", ""))),
        ("Citations", str(draft.get("citation_count", ""))),
    ]

    for label, value in metadata_pairs:
        if not value:
            continue
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"{label}: ")
        run.bold = True
        paragraph.add_run(str(value))

    doc.add_paragraph()

    for section in draft.get("sections", []):
        heading = section.get("label") or section.get("section_name") or section.get("id")
        content = section.get("content") or section.get("text") or ""

        if heading:
            doc.add_heading(str(heading), level=1)

        render_html_to_docx(doc, content)
        doc.add_paragraph()

    if draft.get("provenance"):
        doc.add_heading("Legal Sources & Citations", level=1)
        for source in draft["provenance"]:
            paragraph = doc.add_paragraph()
            title = source.get("source_title", "Unknown Source")
            citation = source.get("citation_format") or source.get("section")
            excerpt = source.get("text_excerpt", "")
            run = paragraph.add_run(f"{title}")
            run.bold = True
            if citation:
                paragraph.add_run(f" ({citation})")
            if excerpt:
                paragraph.add_run(f"\nExcerpt: {excerpt}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def create_docx_response(draft_id: str) -> Response:
    if draft_id not in draft_registry:
        raise HTTPException(status_code=404, detail="Draft not found")

    if not docx_available or Document is None:
        raise HTTPException(
            status_code=503,
            detail=(
                "DOCX export requires python-docx. Install with `pip install python-docx` "
                "or add the dependency to your environment."
            ),
        )

    draft = draft_registry[draft_id]
    docx_bytes = build_petition_docx(draft)

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=petition-{draft_id}.docx",
        },
    )


# -----------------------------------------------------------------------------
# Legacy Compatibility Helpers
# -----------------------------------------------------------------------------


LEGACY_CASE_TYPE_MAP: Dict[str, CaseType] = {
    "civil_revision": CaseType.CIVIL_REVISION,
    "constitutional_writ": CaseType.CONSTITUTIONAL_PETITION,
    "constitutional_petition": CaseType.CONSTITUTIONAL_PETITION,
    "criminal_bail": CaseType.BAIL_APPLICATION,
    "bail_application": CaseType.BAIL_APPLICATION,
    "criminal_revision": CaseType.CRIMINAL_REVISION,
    "civil_appeal": CaseType.CIVIL_APPEAL,
    "criminal_appeal": CaseType.CRIMINAL_APPEAL,
    "civil_suit": CaseType.CIVIL_SUIT,
    "service_matter": CaseType.SERVICE_MATTER,
    "tax_matter": CaseType.TAX_MATTER,
}


LEGACY_COURT_MAP: Dict[str, CourtType] = {
    "supreme court": CourtType.SUPREME_COURT,
    "supreme court of pakistan": CourtType.SUPREME_COURT,
    "lahore high court": CourtType.LAHORE_HIGH_COURT,
    "sindh high court": CourtType.SINDH_HIGH_COURT,
    "islamabad high court": CourtType.ISLAMABAD_HIGH_COURT,
    "peshawar high court": CourtType.PESHAWAR_HIGH_COURT,
    "balochistan high court": CourtType.BALOCHISTAN_HIGH_COURT,
    "sessions court": CourtType.SESSIONS_COURT,
    "district court": CourtType.DISTRICT_COURT,
}


def resolve_case_type(case_type: str) -> CaseType:
    normalized = (case_type or "").strip().lower()
    if normalized in LEGACY_CASE_TYPE_MAP:
        return LEGACY_CASE_TYPE_MAP[normalized]
    try:
        return CaseType(normalized)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=f"Unsupported case_type: {case_type}") from exc


def resolve_court(jurisdiction: str) -> CourtType:
    normalized = (jurisdiction or "").strip().lower()
    if normalized in LEGACY_COURT_MAP:
        return LEGACY_COURT_MAP[normalized]
    for court in CourtType:
        if court.value.lower() == normalized:
            return court
    raise HTTPException(status_code=400, detail=f"Unsupported court/jurisdiction: {jurisdiction}")


def legacy_case_to_petition_request(case_data: LegacyCaseData) -> PetitionRequest:
    case_type_enum = resolve_case_type(case_data.case_type)
    court_enum = resolve_court(case_data.jurisdiction)

    parties_payload = {
        "petitioner": {
            "name": case_data.parties.petitioner or "[Petitioner]",
            "cnic": "",
            "address": "",
        },
        "respondent": {
            "name": case_data.parties.respondent or "[Respondent]",
            "cnic": "",
            "address": "",
        },
    }

    annexures_payload = [
        {"description": item} if isinstance(item, str) else item
        for item in (case_data.annexures or [])
    ]

    relief_text = case_data.prayers or "Appropriate relief in the interest of justice."

    return PetitionRequest(
        case_type=case_type_enum,
        court=court_enum,
        parties=parties_payload,
        facts=case_data.facts,
        specific_provisions_challenged=[],
        relief_sought=relief_text,
        urgency_factors={},
        case_details={},
        annexures=annexures_payload,
    )


def convert_petition_to_legacy(petition: Dict[str, Any]) -> Dict[str, Any]:
    sections = [
        {
            "section_name": section.get("label") or section.get("id", "Section"),
            "content": section.get("content", ""),
        }
        for section in petition.get("sections", [])
    ]

    validation_data = petition.get("validation", {}) or {}
    overall_score = float(validation_data.get("score", 0.0))
    legacy_checks = []
    for check in validation_data.get("checks", []) or []:
        legacy_checks.append(
            {
                "check_name": check.get("check_name", "validation_check"),
                "status": check.get("status", "info"),
                "message": check.get("message", ""),
                "severity": check.get("severity", "info"),
                "fix_suggestion": check.get("fix_suggestion"),
            }
        )

    legacy_validation = {
        "overall_score": overall_score,
        "status": validation_data.get("status", "warn"),
        "checks": legacy_checks,
    }

    coverage_score = petition.get("coverage_score")
    if coverage_score is None:
        coverage_score = overall_score

    return {
        "draft_id": petition.get("draft_id"),
        "template_version": petition.get("template_version"),
        "created_at": petition.get("created_at"),
        "court": petition.get("court"),
        "case_type": petition.get("case_type"),
        "meta": petition.get("meta", {}),
        "sections": sections,
        "validation": legacy_validation,
        "provenance": petition.get("provenance", []),
        "annexures": petition.get("annexures", []),
        "coverage_score": coverage_score,
    }
# -----------------------------------------------------------------------------
# API Endpoints
# -----------------------------------------------------------------------------


@app.get("/")
async def root() -> Dict[str, Any]:
    return {
        "message": "Elite Pakistani Legal Petition AI System",
        "version": "2.0.0",
        "status": "online",
    }


@app.get("/api/v1/health")
async def legacy_health_check() -> Dict[str, Any]:
    """Backward-compatible health endpoint for v1 frontend."""
    status = await health_check()
    return status


@app.get("/api/v1/templates")
async def legacy_templates() -> Dict[str, Any]:
    """Expose template list for legacy frontend."""

    templates = [
        {
            "id": key,
            "name": value["meta"].get("case_type", key.replace("_", " ").title()),
            "description": ", ".join(value["meta"].get("legal_basis", [])),
        }
        for key, value in COMPREHENSIVE_TEMPLATES.items()
    ]
    return {"templates": templates}


@app.post("/api/v1/petitions/generate")
async def generate_petition_v1(
    case_data: LegacyCaseData, background: BackgroundTasks
) -> JSONResponse:
    await ensure_services_ready()

    petition_request = legacy_case_to_petition_request(case_data)

    try:
        petition = await ElitePetitionGenerator.generate_complete_petition(petition_request)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except HTTPException:
        raise
    except Exception as exc:  # pragma: no cover
        logger.exception("Legacy petition generation failed: %s", exc)
        raise HTTPException(status_code=500, detail="Failed to generate petition") from exc

    draft_registry[petition["draft_id"]] = petition
    legacy_payload = convert_petition_to_legacy(petition)
    background.add_task(logger.info, "Legacy petition %s generated", petition["draft_id"])
    return JSONResponse(content=legacy_payload)


@app.post("/api/v1/chat")
async def chat_v1(request: LegacyChatRequest) -> Dict[str, Any]:
    if not request.message.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty")

    await ensure_services_ready()

    session_id = request.session_id or str(uuid4())

    context_chunks: List[str] = []
    if vector_store:
        try:
            matches = await vector_store.search_with_filters(
                query=request.message,
                case_type="general",
                court="All Courts",
                top_k=3,
            )
            for match in matches:
                metadata = match.get("metadata", {})
                context_chunks.append(
                    f"Source: {metadata.get('source_title', 'Unknown')}\n"
                    f"Excerpt: {match.get('text', '')[:400]}"
                )
        except Exception as exc:  # pragma: no cover
            logger.warning("Vector retrieval for chat failed: %s", exc)

    context_text = "\n\n".join(context_chunks) if context_chunks else "No specific sources retrieved."

    prompt = (
        "You are an expert Pakistani legal assistant. Provide precise, cited guidance "
        "based on Pakistani law, statutes, and jurisprudence."
    )

    try:
        response = await asyncio.to_thread(
            openai_client.chat.completions.create,
            model="gpt-4o",
            messages=[
                {"role": "system", "content": prompt},
                {
                    "role": "user",
                    "content": (
                        f"Question: {request.message}\n\n"
                        f"Relevant Sources:\n{context_text}\n\n"
                        "Provide a comprehensive answer with legal references."
                    ),
                },
            ],
            temperature=0.4,
            max_tokens=800,
        )
        content = response.choices[0].message.content
    except Exception as exc:
        logger.error("Chat completion failed: %s", exc)
        raise HTTPException(status_code=502, detail="Failed to generate chat response") from exc

    return {"response": content, "session_id": session_id}


@app.post("/api/v1/petitions/{draft_id}/finalize")
async def finalize_petition_v1(draft_id: str, request: LegacyFinalizeRequest) -> Dict[str, Any]:
    draft = draft_registry.get(draft_id)
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    
    draft["status"] = "finalized"
    draft["finalized_at"] = datetime.now().isoformat()
    draft["approver_name"] = request.approver_name
    draft["approver_id"] = request.approver_id
    draft["notes"] = request.notes
    
    return {
        "draft_id": draft_id,
        "status": "finalized",
        "message": "Petition finalized successfully",
    }


@app.get("/api/v1/petitions/{draft_id}/docx")
async def download_docx_v1(draft_id: str) -> Response:
    return create_docx_response(draft_id)


@app.get("/api/v1/petitions/{draft_id}/pdf")
async def deprecated_pdf_v1(draft_id: str) -> Response:
    raise HTTPException(
        status_code=410,
        detail="PDF export is deprecated. Please use /api/v1/petitions/{draft_id}/docx.",
    )


@app.get("/api/v2/health")
async def health_check() -> Dict[str, Any]:
    openai_status = "connected" if openai_client else "not_configured"
    pinecone_status = (
        "connected" if vector_store and vector_store.index is not None else "not_configured"
    )
    return {
        "timestamp": datetime.now().isoformat(),
        "openai": openai_status,
        "pinecone": pinecone_status,
        "status": "healthy" if openai_status == "connected" else "degraded",
    }


@app.post("/api/v2/petitions/generate")
async def generate_petition_endpoint(request: PetitionRequest, background: BackgroundTasks) -> JSONResponse:
    await ensure_services_ready()

    try:
        petition = await ElitePetitionGenerator.generate_complete_petition(request)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except HTTPException:
        raise
    except Exception as exc:  # pragma: no cover
        logger.exception("Unexpected generation error: %s", exc)
        raise HTTPException(status_code=500, detail="Failed to generate petition") from exc

    draft_registry[petition["draft_id"]] = petition
    background.add_task(logger.info, "Petition %s generated", petition["draft_id"])
    return JSONResponse(content=petition)


@app.post("/api/v2/petitions/generate/stream")
async def generate_petition_stream(request: PetitionRequest) -> StreamingResponse:
    await ensure_services_ready()

    async def event_stream() -> Any:
        try:
            petition = await ElitePetitionGenerator.generate_complete_petition(request)
            draft_registry[petition["draft_id"]] = petition
            for section in petition["sections"]:
                payload = json.dumps({"event": "section", "data": section})
                yield payload + "\n"
                await asyncio.sleep(0)
            yield json.dumps({"event": "complete", "data": petition}) + "\n"
        except Exception as exc:  # pragma: no cover
            logger.error("Streaming generation error: %s", exc)
            yield json.dumps({"event": "error", "message": str(exc)}) + "\n"

    return StreamingResponse(event_stream(), media_type="application/jsonl")


@app.get("/api/v2/petitions/{draft_id}")
async def get_petition(draft_id: str) -> JSONResponse:
    draft = draft_registry.get(draft_id)
    if not draft:
        raise HTTPException(status_code=404, detail="Draft not found")
    return JSONResponse(content=draft)


@app.get("/api/v2/petitions/{draft_id}/docx")
async def download_docx(draft_id: str) -> Response:
    return create_docx_response(draft_id)


@app.get("/api/v2/petitions/{draft_id}/pdf")
async def deprecated_pdf(draft_id: str) -> Response:
    raise HTTPException(
        status_code=410,
        detail="PDF export is deprecated. Please use /api/v2/petitions/{draft_id}/docx.",
    )


@app.delete("/api/v2/petitions/{draft_id}")
async def delete_petition(draft_id: str) -> Dict[str, str]:
    if draft_id in draft_registry:
        del draft_registry[draft_id]
        return {"message": "Draft deleted"}
    raise HTTPException(status_code=404, detail="Draft not found")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)

